"""
議会議事録管理システム（Streamlit版）
"""

import streamlit as st
import pandas as pd
from datetime import datetime
import sqlite3
import os
from pathlib import Path
import docx
from io import BytesIO

# ページ設定
st.set_page_config(
    page_title="議会議事録管理システム",
    page_icon="📝",
    layout="wide",
    initial_sidebar_state="expanded"
)

# データベース接続
@st.cache_resource
def get_db_connection():
    conn = sqlite3.connect('transcription.db', check_same_thread=False)
    conn.row_factory = sqlite3.Row
    return conn

# 初期化
def init_db():
    conn = get_db_connection()
    conn.execute('''
        CREATE TABLE IF NOT EXISTS speakers (
            id INTEGER PRIMARY KEY,
            full_name TEXT NOT NULL,
            display_name TEXT,
            aliases TEXT,
            speaker_type TEXT
        )
    ''')
    conn.execute('''
        CREATE TABLE IF NOT EXISTS sessions (
            id INTEGER PRIMARY KEY,
            name TEXT NOT NULL,
            date TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    conn.execute('''
        CREATE TABLE IF NOT EXISTS sections (
            id INTEGER PRIMARY KEY,
            session_id INTEGER,
            speaker_id INTEGER,
            speaker_name TEXT,
            content TEXT,
            timestamp TEXT,
            is_included BOOLEAN DEFAULT 1,
            FOREIGN KEY (session_id) REFERENCES sessions(id),
            FOREIGN KEY (speaker_id) REFERENCES speakers(id)
        )
    ''')
    conn.commit()

# サイドバー
with st.sidebar:
    st.title("📝 議事録管理")
    
    page = st.radio(
        "メニュー",
        ["🏠 ホーム", "📂 セッション管理", "👥 話者管理", "📊 分析"]
    )

# メインコンテンツ
if page == "🏠 ホーム":
    st.title("議会議事録管理システム")
    st.markdown("""
    ### 主な機能
    - 📁 議事録データのアップロード（NOTTA/MANUS形式）
    - 👥 話者の登録・管理
    - ⏱️ 話者別積算時間の計算
    - 📄 Word形式での出力
    """)
    
    # クイックアクション
    col1, col2, col3 = st.columns(3)
    with col1:
        if st.button("📂 新規セッション作成", use_container_width=True):
            st.session_state.page = "session"
    with col2:
        if st.button("👥 話者を登録", use_container_width=True):
            st.session_state.page = "speaker"
    with col3:
        if st.button("📊 統計を見る", use_container_width=True):
            st.session_state.page = "analytics"

elif page == "📂 セッション管理":
    st.title("セッション管理")
    
    tab1, tab2, tab3 = st.tabs(["セッション一覧", "新規作成", "編集"])
    
    with tab1:
        conn = get_db_connection()
        sessions = pd.read_sql_query(
            "SELECT * FROM sessions ORDER BY date DESC", 
            conn
        )
        
        if not sessions.empty:
            st.dataframe(sessions, use_container_width=True)
            
            # セッション選択
            selected_id = st.selectbox(
                "編集するセッションを選択",
                sessions['id'].tolist(),
                format_func=lambda x: sessions[sessions['id']==x]['name'].iloc[0]
            )
            
            if st.button("編集"):
                st.session_state.editing_session = selected_id
                st.rerun()
        else:
            st.info("セッションがありません")
    
    with tab2:
        with st.form("new_session"):
            st.subheader("新規セッション作成")
            name = st.text_input("セッション名")
            date = st.date_input("日付")
            
            uploaded_file = st.file_uploader(
                "議事録ファイル（TXT形式）",
                type=['txt']
            )
            
            if st.form_submit_button("作成"):
                if name and uploaded_file:
                    # セッション作成
                    conn = get_db_connection()
                    cursor = conn.cursor()
                    cursor.execute(
                        "INSERT INTO sessions (name, date) VALUES (?, ?)",
                        (name, date.isoformat())
                    )
                    session_id = cursor.lastrowid
                    
                    # ファイル処理
                    content = uploaded_file.read().decode('utf-8')
                    sections = parse_transcription(content)
                    
                    # セクション保存
                    for section in sections:
                        cursor.execute("""
                            INSERT INTO sections 
                            (session_id, speaker_name, content, timestamp)
                            VALUES (?, ?, ?, ?)
                        """, (
                            session_id,
                            section['speaker'],
                            section['content'],
                            section['timestamp']
                        ))
                    
                    conn.commit()
                    st.success(f"セッション '{name}' を作成しました")
                    st.rerun()
    
    with tab3:
        if 'editing_session' in st.session_state:
            session_id = st.session_state.editing_session
            conn = get_db_connection()
            
            # セクション取得
            sections = pd.read_sql_query(
                "SELECT * FROM sections WHERE session_id = ?",
                conn,
                params=(session_id,)
            )
            
            st.subheader("セクション編集")
            
            # チェックボックスで含める/除外を管理
            for idx, row in sections.iterrows():
                col1, col2, col3 = st.columns([1, 3, 6])
                
                with col1:
                    included = st.checkbox(
                        "含める",
                        value=bool(row['is_included']),
                        key=f"inc_{row['id']}"
                    )
                    
                    if included != row['is_included']:
                        conn.execute(
                            "UPDATE sections SET is_included = ? WHERE id = ?",
                            (included, row['id'])
                        )
                        conn.commit()
                
                with col2:
                    st.text(f"{row['speaker_name']}")
                    st.caption(f"{row['timestamp']}")
                
                with col3:
                    st.text_area(
                        "内容",
                        value=row['content'],
                        key=f"content_{row['id']}",
                        height=100,
                        disabled=True
                    )
            
            # Word出力
            if st.button("📄 Word形式でダウンロード"):
                doc = create_word_document(session_id)
                
                bio = BytesIO()
                doc.save(bio)
                bio.seek(0)
                
                st.download_button(
                    label="ダウンロード",
                    data=bio.read(),
                    file_name=f"議事録_{datetime.now().strftime('%Y%m%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

elif page == "👥 話者管理":
    st.title("話者管理")
    
    tab1, tab2 = st.tabs(["話者一覧", "新規登録"])
    
    with tab1:
        conn = get_db_connection()
        speakers = pd.read_sql_query(
            "SELECT * FROM speakers ORDER BY speaker_type, full_name",
            conn
        )
        
        if not speakers.empty:
            # 編集可能なデータフレーム
            edited_df = st.data_editor(
                speakers,
                hide_index=True,
                use_container_width=True,
                column_config={
                    "speaker_type": st.column_config.SelectboxColumn(
                        "種別",
                        options=["MEMBER", "STAFF", "OTHER"]
                    )
                }
            )
            
            # 更新ボタン
            if st.button("更新を保存"):
                # 変更を検出して更新
                for idx, row in edited_df.iterrows():
                    conn.execute("""
                        UPDATE speakers 
                        SET full_name = ?, display_name = ?, 
                            aliases = ?, speaker_type = ?
                        WHERE id = ?
                    """, (
                        row['full_name'], row['display_name'],
                        row['aliases'], row['speaker_type'], row['id']
                    ))
                conn.commit()
                st.success("更新しました")
        else:
            st.info("話者が登録されていません")
    
    with tab2:
        with st.form("new_speaker"):
            st.subheader("新規話者登録")
            
            col1, col2 = st.columns(2)
            with col1:
                full_name = st.text_input("正式名称")
                speaker_type = st.selectbox(
                    "種別",
                    ["MEMBER", "STAFF", "OTHER"],
                    format_func=lambda x: {
                        "MEMBER": "議員",
                        "STAFF": "職員",
                        "OTHER": "その他"
                    }[x]
                )
            
            with col2:
                display_name = st.text_input(
                    "表示名",
                    value=f"{full_name}{'議員' if speaker_type == 'MEMBER' else ''}"
                )
                aliases = st.text_input(
                    "別名（カンマ区切り）",
                    placeholder="例: 比嘉,ヒガ,比嘉議員"
                )
            
            if st.form_submit_button("登録"):
                if full_name:
                    conn = get_db_connection()
                    conn.execute("""
                        INSERT INTO speakers 
                        (full_name, display_name, aliases, speaker_type)
                        VALUES (?, ?, ?, ?)
                    """, (full_name, display_name, aliases, speaker_type))
                    conn.commit()
                    st.success(f"話者 '{full_name}' を登録しました")
                    st.rerun()

elif page == "📊 分析":
    st.title("議事録分析")
    
    conn = get_db_connection()
    
    # 基本統計
    col1, col2, col3 = st.columns(3)
    
    with col1:
        session_count = conn.execute("SELECT COUNT(*) FROM sessions").fetchone()[0]
        st.metric("セッション数", session_count)
    
    with col2:
        speaker_count = conn.execute("SELECT COUNT(*) FROM speakers").fetchone()[0]
        st.metric("登録話者数", speaker_count)
    
    with col3:
        section_count = conn.execute("SELECT COUNT(*) FROM sections").fetchone()[0]
        st.metric("総発言数", section_count)
    
    # 話者別統計
    st.subheader("話者別発言統計")
    
    speaker_stats = pd.read_sql_query("""
        SELECT 
            s.speaker_name,
            COUNT(*) as count,
            SUM(LENGTH(s.content)) as total_chars
        FROM sections s
        WHERE s.is_included = 1
        GROUP BY s.speaker_name
        ORDER BY count DESC
        LIMIT 20
    """, conn)
    
    if not speaker_stats.empty:
        # グラフ表示
        st.bar_chart(speaker_stats.set_index('speaker_name')['count'])
        
        # 詳細表
        st.dataframe(speaker_stats, use_container_width=True)

# ヘルパー関数
def parse_transcription(content: str) -> list:
    """テキストファイルから発言を解析"""
    sections = []
    lines = content.strip().split('\n')
    
    current_section = None
    for line in lines:
        # 簡易的なパーサー（実際の形式に合わせて調整）
        if line.strip() and not line.startswith(' '):
            if current_section:
                sections.append(current_section)
            
            # 話者と時刻の抽出
            parts = line.split('：', 1)
            if len(parts) == 2:
                speaker = parts[0].strip()
                current_section = {
                    'speaker': speaker,
                    'content': parts[1].strip(),
                    'timestamp': '00:00:00'  # 実際には解析が必要
                }
        elif current_section and line.strip():
            current_section['content'] += '\n' + line.strip()
    
    if current_section:
        sections.append(current_section)
    
    return sections

def create_word_document(session_id: int):
    """Word文書を生成"""
    doc = docx.Document()
    conn = get_db_connection()
    
    # セッション情報
    session = conn.execute(
        "SELECT * FROM sessions WHERE id = ?",
        (session_id,)
    ).fetchone()
    
    doc.add_heading(session['name'], 0)
    doc.add_paragraph(f"日付: {session['date']}")
    
    # セクション
    sections = conn.execute("""
        SELECT * FROM sections 
        WHERE session_id = ? AND is_included = 1
        ORDER BY id
    """, (session_id,)).fetchall()
    
    for section in sections:
        p = doc.add_paragraph()
        p.add_run(f"【{section['speaker_name']}】").bold = True
        p.add_run(f" {section['timestamp']}")
        doc.add_paragraph(section['content'])
        doc.add_paragraph()  # 空行
    
    return doc

# 初期化
init_db()